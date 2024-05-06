from docx import Document
from docx2pdf import convert
from pathlib import Path

from docx.shared import Pt  # Import Pt (point) for font size

import os
import asyncio
def get_unique_filename(directory, base_filename, extension):
    count = 1
    filename = os.path.join(directory, base_filename + "." + extension)

    while os.path.exists(filename):
        filename = os.path.join(directory, f"{base_filename}_{count}.{extension}")
        count += 1

    return filename

def get_unique_filenames(directory, base_filename):
    extension_word = "docx"
    extension_pdf = "pdf"
    count = 1

    wordname = os.path.join(directory, base_filename + "." + extension_word)
    pdfname = os.path.join(directory, base_filename + "." + extension_pdf)

    while os.path.exists(wordname):
        wordname = os.path.join(directory, f"{base_filename}_{count}.{extension_word}")
        pdfname = os.path.join(directory, f"{base_filename}_{count}.{extension_pdf}")
        count += 1

    return (wordname, pdfname)

will_converted = []

def upload(job, skills, case):
    # Load the existing resume from the Word document
    filename = job["company"]
    description = job["description"]
    if case == "Data":
        doc = Document("./data/new-resume-format-data.docx")
    elif case == "Backend":
        doc = Document("./data/new-resume-format-backend.docx")

    # Specify the path of the folder you want to create
    folder_path = Path('./data/company/' + filename)

    # Use the mkdir() method to create the folder
    if(not folder_path.exists()):
        folder_path.mkdir()

    synonyms = {
        "Google Cloud Platform": "GCP (Google Cloud Platform)",
        "Google Cloud Platform (GCP)": "GCP (Google Cloud Platform)",
        "Google Cloud Platform(GCP)": "GCP (Google Cloud Platform)",
        "GCP": "GCP (Google Cloud Platform)",
        "MongoDb": "MongoDB",
        "Programming Interface (API)": "API",
        "Programming Interface": "API",
        "(API) Programming Interface": "API",
        "Object Oriented Programming" :"(OOP) Object Oriented Programming",
        "OOP" :"(OOP) Object Oriented Programming",
        "JSON Web Tokens": "JWT",
        "Cloud": "Cloud Computing",
        "Apache Druid": "Druid",
        "Big Query": "BigQuery",
        "Java/Scala": "Java",
        "REST": "RESTFUL Web Services"
    }

    # Normalize words
    normalized_skills = [synonyms.get(word.strip(), word) for word in skills]

    # Remove duplicates
    unique_words = list(set(normalized_skills))

    # Sort the list
    unique_words.sort()

    # Update the technical skills section
    skills_text = "\n".join(unique_words).strip()

    # Find the technical skills section and insert new skills
    found_technical_skills = False
    names = get_unique_filenames("./data/company/" + filename + "/", "Resume-Utku-Aysev")
    docx_name = names[0]
    pdf_name = names[1]
    # Save the updated resume to a new Word document
    doc.save(docx_name)
    will_converted.append((docx_name, pdf_name))

    doc = Document()
    # Add content to the document (e.g., description)
    doc.add_heading('Cover Letter', 0)  # Add a heading for the cover letter
    doc.add_paragraph(description)      # Add the description to the document
    # Save the document with the specified filename
    doc.save(docx_name.replace("Resume-Utku-Aysev", "Cover-Letter"))

    if len(unique_words) > 0:
        with open(docx_name.replace("Resume-Utku-Aysev", "Missing-Skills").replace("docx","txt"), 'w') as file:
            file.write(skills_text)


def convert_all():
    for elem in will_converted:
        print(elem[0] + " " + elem[1])
        convert(elem[0], elem[1])

